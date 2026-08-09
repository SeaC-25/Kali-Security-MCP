#!/usr/bin/env python3
"""Task report export: Markdown (primary) + structured summary paths.

Product rules:
- Main table: verified findings only (raw truth, no redaction by default)
- Candidates / insight hypotheses go to appendix
- Timeline from action_log; ATT&CK coverage as label appendix only
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Dict, List, Optional

from kali_mcp.core.action_log import read_actions, task_timeline
from kali_mcp.core.attack_coverage import task_attack_coverage
from kali_mcp.core.findings_store import list_findings
from kali_mcp.core.handoff import compile_handoff
from kali_mcp.core.observer import analyze_actions
from kali_mcp.core.target_graph import get_graph
from kali_mcp.core.task_workspace import get_workspace, utc_now_iso


def _esc_md(text: Any) -> str:
    s = "" if text is None else str(text)
    return s.replace("\r\n", "\n").replace("\r", "\n")


def _bullet_list(items: List[str], empty: str = "- （无）") -> str:
    cleaned = [x for x in items if x]
    if not cleaned:
        return empty
    return "\n".join(f"- {x}" for x in cleaned)


def _finding_block(f: Dict[str, Any], idx: int) -> str:
    fid = f.get("finding_id") or ""
    title = f.get("title") or "(untitled)"
    sev = f.get("severity") or "info"
    status = f.get("status") or "candidate"
    target = f.get("target") or f.get("asset") or ""
    source = f.get("source") or ""
    cmd = f.get("reproduce_cmd") or ""
    signal = f.get("expected_signal") or ""
    tactics = f.get("tactic") or ""
    techs = f.get("technique_ids") or []
    if not isinstance(techs, list):
        techs = []
    evidence = f.get("evidence_paths") or []
    if not isinstance(evidence, list):
        evidence = []
    notes = f.get("raw_notes") or f.get("notes") or ""
    lines = [
        f"### {idx}. {title}",
        "",
        f"- **finding_id**: `{fid}`",
        f"- **status**: {status}",
        f"- **severity**: {sev}",
        f"- **target**: `{target}`",
        f"- **source**: {source}",
    ]
    if tactics:
        lines.append(f"- **tactic**: {tactics}")
    if techs:
        lines.append(f"- **technique_ids**: {', '.join(str(t) for t in techs)}")
    if cmd:
        lines.extend(["", "**复现命令（真值）**", "", "```", _esc_md(cmd), "```"])
    if signal:
        lines.append(f"- **expected_signal**: `{signal}`")
    if evidence:
        lines.append("- **evidence_paths**:")
        for p in evidence:
            lines.append(f"  - `{p}`")
    if notes:
        lines.extend(["", "**raw_notes**", "", "```", _esc_md(notes), "```"])
    lines.append("")
    return "\n".join(lines)


def build_report_payload(task_id: str) -> Dict[str, Any]:
    """Assemble structured report data (no disk write)."""
    if not task_id:
        return {"ok": False, "error": "task_id required"}
    ws = get_workspace(task_id, create=True)
    graph = get_graph(ws.task_id)
    findings = list_findings(ws.task_id)
    verified = [f for f in findings if f.get("status") == "verified"]
    candidates = [f for f in findings if f.get("status") == "candidate"]
    other = [f for f in findings if f.get("status") not in ("verified", "candidate")]
    actions = read_actions(ws.task_id, limit=500)
    timeline = task_timeline(ws.task_id, limit=500)
    coverage = task_attack_coverage(ws.task_id)
    observer = analyze_actions(ws.task_id)
    handoff = compile_handoff(ws.task_id)
    meta = ws.read_meta()
    return {
        "ok": True,
        "task_id": ws.task_id,
        "exported_at": utc_now_iso(),
        "meta": meta,
        "graph_summary": graph.summary(),
        "findings": findings,
        "verified": verified,
        "candidates": candidates,
        "other_findings": other,
        "actions": actions,
        "timeline": {
            "count": timeline.get("count"),
            "duration_ms_sum": timeline.get("duration_ms_sum"),
            "duration_ms_known_count": timeline.get("duration_ms_known_count"),
            "by_tool": timeline.get("by_tool"),
            "by_source": timeline.get("by_source"),
        },
        "attack_coverage": coverage,
        "observer": {
            "duplicate_count": observer.get("duplicate_count"),
            "suggestions": observer.get("suggestions"),
            "should_slow_down": observer.get("should_slow_down"),
        },
        "handoff": {
            "handoff_json": handoff.get("handoff_json"),
            "progress_md": handoff.get("progress_md"),
        },
        "workspace_root": str(ws.root),
    }


def _docx_available() -> bool:
    try:
        import docx  # noqa: F401

        return True
    except ImportError:
        return False


def render_docx(payload: Dict[str, Any], path: Optional[Path] = None) -> Dict[str, Any]:
    """Build Word report aligned with markdown sections. Requires python-docx."""
    if not payload.get("ok"):
        return {"ok": False, "error": payload.get("error") or "payload not ok"}
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        return {
            "ok": False,
            "error": "python-docx not installed; pip install python-docx",
            "missing_dependency": "python-docx",
        }

    tid = payload.get("task_id") or ""
    meta = payload.get("meta") or {}
    graph = payload.get("graph_summary") or {}
    verified: List[Dict[str, Any]] = list(payload.get("verified") or [])
    candidates: List[Dict[str, Any]] = list(payload.get("candidates") or [])
    other: List[Dict[str, Any]] = list(payload.get("other_findings") or [])
    actions: List[Dict[str, Any]] = list(payload.get("actions") or [])
    cov = payload.get("attack_coverage") or {}
    cov_all = cov.get("all") or {}
    cov_v = cov.get("verified_only") or {}
    observer = payload.get("observer") or {}
    handoff = payload.get("handoff") or {}
    targets = meta.get("targets") or []
    if isinstance(targets, str):
        targets = [targets]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(f"渗透任务报告 — {tid}", level=0)
    meta_lines = [
        f"导出时间: {payload.get('exported_at') or ''}",
        f"任务状态: {meta.get('status') or ''}",
        f"阶段: {meta.get('phase') or ''}",
        f"深度: {meta.get('depth') or ''}",
        f"目标: {', '.join(str(t) for t in targets) if targets else '（无）'}",
        f"工作区: {payload.get('workspace_root') or ''}",
        "口径：主表仅 verified；证据与复现命令默认不脱敏（真值）。",
    ]
    for line in meta_lines:
        doc.add_paragraph(line)

    doc.add_heading("1. 摘要", level=1)
    for line in (
        f"图节点: {graph.get('nodes', 0)} / 边: {graph.get('edges', 0)} / dead: {graph.get('dead_nodes', 0)}",
        f"verified findings: {len(verified)}",
        f"candidate findings: {len(candidates)}",
        f"other (fp/blocked/...): {len(other)}",
        f"action log 条数（最近窗口）: {len(actions)}",
        f"Observer 重复建议: {observer.get('duplicate_count', 0)}",
    ):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("2. Verified Findings（主表）", level=1)
    if not verified:
        doc.add_paragraph("（无 verified；若仅有 candidate 请先 verify_finding）")
    else:
        for i, f in enumerate(verified, 1):
            title = f.get("title") or "(untitled)"
            doc.add_heading(f"{i}. {title}", level=2)
            for key, label in (
                ("finding_id", "finding_id"),
                ("status", "status"),
                ("severity", "severity"),
                ("target", "target"),
                ("source", "source"),
                ("tactic", "tactic"),
                ("expected_signal", "expected_signal"),
            ):
                val = f.get(key)
                if val:
                    doc.add_paragraph(f"{label}: {val}")
            techs = f.get("technique_ids") or []
            if isinstance(techs, list) and techs:
                doc.add_paragraph(f"technique_ids: {', '.join(str(t) for t in techs)}")
            cmd = f.get("reproduce_cmd") or ""
            if cmd:
                doc.add_paragraph("复现命令（真值）:")
                doc.add_paragraph(str(cmd))
            evidence = f.get("evidence_paths") or []
            if isinstance(evidence, list) and evidence:
                doc.add_paragraph("evidence_paths:")
                for p in evidence:
                    doc.add_paragraph(str(p), style="List Bullet")
            notes = f.get("raw_notes") or f.get("notes") or ""
            if notes:
                doc.add_paragraph("raw_notes:")
                doc.add_paragraph(str(notes))

    timeline = payload.get("timeline") or {}
    dur_sum = timeline.get("duration_ms_sum")
    if dur_sum is None:
        dur_sum = 0.0
        for a in actions:
            d = a.get("duration_ms")
            if isinstance(d, (int, float)):
                dur_sum += float(d)
        dur_sum = round(float(dur_sum), 3)
    doc.add_heading("3. 执行时间线（action_log）", level=1)
    doc.add_paragraph(f"duration_ms_sum: {dur_sum}")
    if not actions:
        doc.add_paragraph("（无 action）")
    else:
        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, name in enumerate(("#", "ts", "phase", "tool", "target", "exit", "duration_ms", "source")):
            hdr[i].text = name
        for i, a in enumerate(actions, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = str(a.get("ts") or "")[:40]
            row[2].text = str(a.get("phase") or "")
            row[3].text = str(a.get("tool") or "")
            row[4].text = str(a.get("target") or "")[:80]
            row[5].text = str(a.get("exit") if a.get("exit") is not None else "")
            dur = a.get("duration_ms")
            row[6].text = "" if dur is None else str(dur)
            row[7].text = str(a.get("source") or "")

    doc.add_heading("4. ATT&CK 覆盖附录（标签层）", level=1)
    doc.add_heading("全部 findings", level=2)
    doc.add_paragraph(
        f"labeled: {cov_all.get('labeled_count', 0)} / total: {cov_all.get('findings_total', 0)} "
        f"(ratio={cov_all.get('labeled_ratio', 0)})"
    )
    doc.add_paragraph(f"techniques: {', '.join(cov_all.get('techniques') or []) or '（无）'}")
    doc.add_paragraph(f"tactics_present: {', '.join(cov_all.get('tactics_present') or []) or '（无）'}")
    doc.add_heading("verified only", level=2)
    doc.add_paragraph(
        f"labeled: {cov_v.get('labeled_count', 0)} / total: {cov_v.get('findings_total', 0)}"
    )
    doc.add_paragraph(f"techniques: {', '.join(cov_v.get('techniques') or []) or '（无）'}")
    doc.add_paragraph(f"tactics_present: {', '.join(cov_v.get('tactics_present') or []) or '（无）'}")
    doc.add_paragraph(str(cov_all.get("note") or "ATT&CK is label-only"))

    doc.add_heading("5. Candidate / 未证实假设（附录）", level=1)
    if not candidates:
        doc.add_paragraph("（无 candidate）")
    else:
        for i, f in enumerate(candidates, 1):
            title = f.get("title") or "(untitled)"
            doc.add_heading(f"{i}. {title}", level=2)
            doc.add_paragraph(f"status: {f.get('status') or 'candidate'}")
            doc.add_paragraph(f"target: {f.get('target') or ''}")
            doc.add_paragraph(f"source: {f.get('source') or ''}")
            cmd = f.get("reproduce_cmd") or ""
            if cmd:
                doc.add_paragraph(f"reproduce_cmd: {cmd}")

    if other:
        doc.add_heading("6. 其他状态 findings", level=1)
        for i, f in enumerate(other, 1):
            doc.add_paragraph(f"{i}. {f.get('title') or '(untitled)'} [{f.get('status') or ''}]")

    suggestions = observer.get("suggestions") or []
    doc.add_heading("7. Observer 建议（旁路，默认不拦截）", level=1)
    if not suggestions:
        doc.add_paragraph("（无）")
    else:
        for s in suggestions:
            if s is None:
                continue
            if isinstance(s, dict):
                doc.add_paragraph(str(s.get("message") or s), style="List Bullet")
            else:
                doc.add_paragraph(str(s), style="List Bullet")

    doc.add_heading("8. Handoff / 续跑指针", level=1)
    doc.add_paragraph(f"handoff_json: {handoff.get('handoff_json') or ''}")
    doc.add_paragraph(f"progress_md: {handoff.get('progress_md') or ''}")
    doc.add_paragraph(f"generated by kali_mcp.core.report_export @ {payload.get('exported_at') or ''}")

    out: Dict[str, Any] = {"ok": True, "document": doc, "task_id": tid}
    if path is not None:
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
        out["path"] = str(path)
        try:
            out["bytes"] = path.stat().st_size
        except OSError:
            out["bytes"] = 0
    return out


def render_markdown(payload: Dict[str, Any]) -> str:
    """Render full Markdown report from build_report_payload output."""
    if not payload.get("ok"):
        return f"# Report error\n\n{_esc_md(payload.get('error') or 'unknown')}\n"

    tid = payload.get("task_id") or ""
    meta = payload.get("meta") or {}
    graph = payload.get("graph_summary") or {}
    verified: List[Dict[str, Any]] = list(payload.get("verified") or [])
    candidates: List[Dict[str, Any]] = list(payload.get("candidates") or [])
    other: List[Dict[str, Any]] = list(payload.get("other_findings") or [])
    actions: List[Dict[str, Any]] = list(payload.get("actions") or [])
    timeline = payload.get("timeline") or {}
    cov = payload.get("attack_coverage") or {}
    cov_all = cov.get("all") or {}
    cov_v = cov.get("verified_only") or {}
    observer = payload.get("observer") or {}
    handoff = payload.get("handoff") or {}
    targets = meta.get("targets") or []
    if isinstance(targets, str):
        targets = [targets]
    dur_sum = timeline.get("duration_ms_sum")
    if dur_sum is None:
        dur_sum = 0.0
        for a in actions:
            d = a.get("duration_ms")
            if isinstance(d, (int, float)):
                dur_sum += float(d)
        dur_sum = round(float(dur_sum), 3)
    dur_known = timeline.get("duration_ms_known_count")
    if dur_known is None:
        dur_known = sum(1 for a in actions if isinstance(a.get("duration_ms"), (int, float)))

    parts: List[str] = [
        f"# 渗透任务报告 — `{tid}`",
        "",
        f"- **导出时间**: {payload.get('exported_at') or ''}",
        f"- **任务状态**: {meta.get('status') or ''}",
        f"- **阶段**: {meta.get('phase') or ''}",
        f"- **深度**: {meta.get('depth') or ''}",
        f"- **目标**: {', '.join(f'`{t}`' for t in targets) if targets else '（无）'}",
        f"- **工作区**: `{payload.get('workspace_root') or ''}`",
        "",
        "> 口径：报告主表仅含 **verified**；证据与复现命令默认**不脱敏**（真值）。",
        "",
        "## 1. 摘要",
        "",
        f"- 图节点: **{graph.get('nodes', 0)}** / 边: **{graph.get('edges', 0)}** / dead: {graph.get('dead_nodes', 0)}",
        f"- verified findings: **{len(verified)}**",
        f"- candidate findings: {len(candidates)}",
        f"- other (fp/blocked/...): {len(other)}",
        f"- action log 条数（最近窗口）: {len(actions)}",
        f"- duration_ms_sum: **{dur_sum}**（known={dur_known}）",
        f"- Observer 重复建议: {observer.get('duplicate_count', 0)}",
        "",
        "## 2. Verified Findings（主表）",
        "",
    ]

    if not verified:
        parts.append("_（无 verified；若仅有 candidate 请先 `verify_finding`）_\n")
    else:
        for i, f in enumerate(verified, 1):
            parts.append(_finding_block(f, i))

    parts.extend(
        [
            "## 3. 执行时间线（action_log）",
            "",
            f"- duration_ms_sum: **{dur_sum}** / known events: {dur_known}",
            "",
            "| # | ts | phase | tool | target | exit | duration_ms | source | evidence |",
            "|---|----|-------|------|--------|------|-------------|--------|----------|",
        ]
    )
    if not actions:
        parts.append("| — | — | — | — | — | — | — | — | — |")
    else:
        for i, a in enumerate(actions, 1):
            ts = str(a.get("ts") or "").replace("|", "\\|")
            phase = str(a.get("phase") or "").replace("|", "\\|")
            tool = str(a.get("tool") or "").replace("|", "\\|")
            target = str(a.get("target") or "")[:80].replace("|", "\\|")
            exit_c = str(a.get("exit") if a.get("exit") is not None else "")
            dur = a.get("duration_ms")
            dur_s = "" if dur is None else str(dur)
            source = str(a.get("source") or "").replace("|", "\\|")
            ev = str(a.get("evidence_path") or "")[:60].replace("|", "\\|")
            parts.append(
                f"| {i} | {ts} | {phase} | {tool} | {target} | {exit_c} | {dur_s} | {source} | {ev} |"
            )
    parts.append("")

    parts.extend(
        [
            "## 4. ATT&CK 覆盖附录（标签层）",
            "",
            "### 全部 findings",
            "",
            f"- labeled: {cov_all.get('labeled_count', 0)} / total: {cov_all.get('findings_total', 0)} "
            f"(ratio={cov_all.get('labeled_ratio', 0)})",
            f"- techniques: {', '.join(cov_all.get('techniques') or []) or '（无）'}",
            f"- tactics_present: {', '.join(cov_all.get('tactics_present') or []) or '（无）'}",
            "",
            "### verified only",
            "",
            f"- labeled: {cov_v.get('labeled_count', 0)} / total: {cov_v.get('findings_total', 0)}",
            f"- techniques: {', '.join(cov_v.get('techniques') or []) or '（无）'}",
            f"- tactics_present: {', '.join(cov_v.get('tactics_present') or []) or '（无）'}",
            "",
            f"_{cov_all.get('note') or 'ATT&CK is label-only'}_",
            "",
            "## 5. Candidate / 未证实假设（附录，不与主表混排）",
            "",
        ]
    )
    if not candidates:
        parts.append("_（无 candidate）_\n")
    else:
        for i, f in enumerate(candidates, 1):
            parts.append(_finding_block(f, i))

    if other:
        parts.extend(["## 6. 其他状态 findings（false_positive / blocked 等）", ""])
        for i, f in enumerate(other, 1):
            parts.append(_finding_block(f, i))

    suggestions = observer.get("suggestions") or []
    suggestion_lines: List[str] = []
    for s in suggestions:
        if s is None:
            continue
        if isinstance(s, dict):
            suggestion_lines.append(str(s.get("message") or s))
        else:
            suggestion_lines.append(str(s))
    parts.extend(
        [
            "## 7. Observer 建议（旁路，默认不拦截）",
            "",
            _bullet_list(suggestion_lines),
            "",
            "## 8. Handoff / 续跑指针",
            "",
            f"- handoff_json: `{handoff.get('handoff_json') or ''}`",
            f"- progress_md: `{handoff.get('progress_md') or ''}`",
            "",
            "---",
            "",
            f"_generated by kali_mcp.core.report_export @ {payload.get('exported_at') or ''}_",
            "",
        ]
    )
    return "\n".join(parts)


def export_task_report(
    task_id: str,
    *,
    formats: Optional[List[str]] = None,
) -> Dict[str, Any]:
    """Write report files under task report/ directory.

    formats: subset of {"json", "markdown", "docx"} (default json+markdown).
    Markdown remains primary; docx is optional (needs python-docx).
    """
    payload = build_report_payload(task_id)
    if not payload.get("ok"):
        return payload

    fmt_set = {str(x).strip().lower() for x in (formats or ["json", "markdown"]) if str(x).strip()}
    if not fmt_set:
        fmt_set = {"json", "markdown"}

    ws = get_workspace(task_id, create=True)
    report_dir: Path = ws.report_dir
    report_dir.mkdir(parents=True, exist_ok=True)

    paths: Dict[str, str] = {}
    docx_note = ""
    docx_path = ""
    if "json" in fmt_set:
        # slim JSON for machines (same shape as historical summary + export meta)
        json_body = {
            "task_id": payload["task_id"],
            "exported_at": payload["exported_at"],
            "meta": payload["meta"],
            "graph_summary": payload["graph_summary"],
            "findings": payload["findings"],
            "verified": payload["verified"],
            "candidates": payload["candidates"],
            "other_findings": payload["other_findings"],
            "recent_actions": payload["actions"][-100:],
            "attack_coverage": payload["attack_coverage"],
            "observer": payload["observer"],
            "handoff": payload["handoff"],
        }
        json_path = report_dir / "summary.json"
        json_path.write_text(json.dumps(json_body, ensure_ascii=False, indent=2), encoding="utf-8")
        paths["json"] = str(json_path)

    md_text = ""
    if "markdown" in fmt_set or "md" in fmt_set:
        md_text = render_markdown(payload)
        md_path = report_dir / "report.md"
        md_path.write_text(md_text, encoding="utf-8")
        paths["markdown"] = str(md_path)

    if "docx" in fmt_set:
        docx_file = report_dir / "report.docx"
        built = render_docx(payload, path=docx_file)
        if built.get("ok"):
            docx_path = str(built.get("path") or docx_file)
            paths["docx"] = docx_path
        else:
            docx_note = str(built.get("error") or "docx export failed")
            if len(fmt_set) == 1:
                return {
                    "ok": False,
                    "error": docx_note,
                    "missing_dependency": built.get("missing_dependency"),
                    "task_id": ws.task_id,
                    "paths": paths,
                    "docx_path": "",
                    "docx_note": docx_note,
                }

    return {
        "ok": True,
        "task_id": ws.task_id,
        "exported_at": payload["exported_at"],
        "paths": paths,
        "report_path": paths.get("json") or paths.get("markdown") or paths.get("docx") or "",
        "markdown_path": paths.get("markdown") or "",
        "docx_path": docx_path or paths.get("docx") or "",
        "verified_count": len(payload.get("verified") or []),
        "candidate_count": len(payload.get("candidates") or []),
        "graph_summary": payload.get("graph_summary"),
        "attack_coverage": payload.get("attack_coverage"),
        "observer_suggestions": len((payload.get("observer") or {}).get("suggestions") or []),
        "handoff_json": (payload.get("handoff") or {}).get("handoff_json"),
        "progress_md": (payload.get("handoff") or {}).get("progress_md"),
        "docx_note": docx_note,
        "markdown_chars": len(md_text),
    }
